"""
Deterministic IEEE docx builder.

Opens the patched template (to inherit styles + page layout), clears the body,
then rebuilds all content from the DocumentModel.  No AI. No inference.
Every formatting decision is encoded in the DocumentModel JSON.
"""
from __future__ import annotations

import copy
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from engine.renderer import styles as S
from engine.schema import (
    Block,
    DocumentMetadata,
    DocumentModel,
    EquationBlock,
    FigureBlock,
    HeadingBlock,
    ListBlock,
    ParagraphBlock,
    Reference,
    TableBlock,
)

TEMPLATE_PATH = Path(__file__).parent.parent.parent.parent / "templates" / "ieee_template.docx"


def _fit_picture_size(img_path: Path, max_width_in: float, max_height_in: float) -> tuple[float, float] | None:
    """Compute (width_in, height_in) that fits the image inside the given box while
    preserving its aspect ratio. Returns None if the image can't be read (caller
    should fall back to width-only sizing)."""
    try:
        with Image.open(img_path) as img:
            px_w, px_h = img.size
    except Exception:
        return None
    if px_w <= 0 or px_h <= 0:
        return None

    aspect = px_h / px_w
    width_in = max_width_in
    height_in = width_in * aspect
    if height_in > max_height_in:
        height_in = max_height_in
        width_in = height_in / aspect
    return width_in, height_in


def _make_sect_pr(cols: int, continuous: bool = True) -> OxmlElement:
    """Build a w:sectPr element using the template's exact pt-notation values."""
    sp = OxmlElement("w:sectPr")

    if continuous:
        t = OxmlElement("w:type")
        t.set(qn("w:val"), "continuous")
        sp.append(t)

    pg = OxmlElement("w:pgSz")
    pg.set(qn("w:w"), S.PAGE_W_PT)
    pg.set(qn("w:h"), S.PAGE_H_PT)
    pg.set(qn("w:code"), S.PAGE_CODE)
    sp.append(pg)

    pm = OxmlElement("w:pgMar")
    pm.set(qn("w:top"),    S.MARGIN_TOP)
    pm.set(qn("w:right"),  S.MARGIN_RIGHT)
    pm.set(qn("w:bottom"), S.MARGIN_BOTTOM)
    pm.set(qn("w:left"),   S.MARGIN_LEFT)
    pm.set(qn("w:header"), S.MARGIN_HEADER)
    pm.set(qn("w:footer"), S.MARGIN_FOOTER)
    pm.set(qn("w:gutter"), S.MARGIN_GUTTER)
    sp.append(pm)

    cl = OxmlElement("w:cols")
    if cols > 1:
        cl.set(qn("w:num"), str(cols))
        cl.set(qn("w:space"), S.COL_SPACE_2)
    else:
        cl.set(qn("w:space"), S.COL_SPACE_1)
    sp.append(cl)

    dg = OxmlElement("w:docGrid")
    dg.set(qn("w:linePitch"), S.DOC_GRID_PITCH)
    sp.append(dg)

    return sp


class DocxBuilder:
    def __init__(
        self,
        storage_base: Path,
        output_dir: Path,
        math_images: dict[str, Path],
    ) -> None:
        self.storage_base = storage_base
        self.output_dir = output_dir
        self.math_images = math_images   # anchor (e.g. "EQ 1") → png path

        # Open template to inherit all styles
        self.doc = Document(str(TEMPLATE_PATH))
        body = self.doc.element.body

        # Clear body content; replace with a clean 2-col body sectPr
        for child in list(body):
            body.remove(child)
        body.append(_make_sect_pr(2, continuous=False))

        # Section-break state: we start in the 1-col title section
        self._in_two_col = False

    # ── Public entry point ────────────────────────────────────────────────────

    def build(self, model: DocumentModel) -> Document:
        self._render_header(model.metadata)
        self._end_title_section()           # transition: 1-col → 2-col body
        for block in model.blocks:
            self._render_block(block)
        if model.references:
            self._render_references(model.references)
        return self.doc

    # ── Header (title, authors, abstract, keywords) ───────────────────────────

    def _render_header(self, meta: DocumentMetadata) -> None:
        # Title
        p = self.doc.add_paragraph(meta.title, style=S.PAPER_TITLE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Conference (optional subtitle)
        if meta.conference:
            p = self.doc.add_paragraph(meta.conference, style=S.PAPER_SUBTITLE)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Authors (comma-separated on one line, then affiliations below)
        author_names = ", ".join(a.name for a in meta.authors)
        p = self.doc.add_paragraph(author_names, style=S.AUTHOR)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Affiliations (one per unique key)
        seen: set[str] = set()
        for author in meta.authors:
            for ref in author.affiliation_refs:
                if ref in seen:
                    continue
                seen.add(ref)
                aff = next((a for a in meta.affiliations if a.key == ref), None)
                if aff is None:
                    continue
                parts = [aff.institution]
                if aff.department:
                    parts.insert(0, aff.department)
                if aff.city and aff.country:
                    parts.append(f"{aff.city}, {aff.country}")
                elif aff.country:
                    parts.append(aff.country)
                p = self.doc.add_paragraph(", ".join(parts), style=S.AFFILIATION)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if author.email:
                    ep = self.doc.add_paragraph(author.email, style=S.AFFILIATION)
                    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Abstract
        abstract_para = self.doc.add_paragraph(style=S.ABSTRACT)
        abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = abstract_para.add_run("Abstract—")
        run.bold = True
        run.italic = True
        abstract_para.add_run(meta.abstract)

        # Keywords
        kw_para = self.doc.add_paragraph(style=S.KEYWORDS)
        kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = kw_para.add_run("Index Terms—")
        run.bold = True
        run.italic = True
        kw_para.add_run(", ".join(meta.keywords))

    # ── Section break helpers ─────────────────────────────────────────────────

    def _end_title_section(self) -> None:
        """Insert the inline sectPr that terminates the 1-col title area."""
        self._add_section_break_para(cols=1)
        self._in_two_col = True

    def _add_section_break_para(self, cols: int) -> None:
        """Empty paragraph whose pPr/sectPr terminates the current section with `cols` columns."""
        para = self.doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        pPr.append(_make_sect_pr(cols, continuous=True))

    def _enter_wide(self) -> None:
        """Switch from 2-col → 1-col for a full-width element."""
        if self._in_two_col:
            self._add_section_break_para(cols=2)
            self._in_two_col = False

    def _exit_wide(self) -> None:
        """Switch back from 1-col → 2-col after the full-width element."""
        self._add_section_break_para(cols=1)
        self._in_two_col = True

    # ── Block dispatcher ──────────────────────────────────────────────────────

    def _render_block(self, block: Block) -> None:
        match block.type:
            case "paragraph":
                self._render_paragraph(block)       # type: ignore[arg-type]
            case "heading":
                self._render_heading(block)          # type: ignore[arg-type]
            case "figure":
                self._render_figure(block)           # type: ignore[arg-type]
            case "wide_figure":
                self._render_wide_figure(block)      # type: ignore[arg-type]
            case "table":
                self._render_table(block)            # type: ignore[arg-type]
            case "wide_table":
                self._render_wide_table(block)       # type: ignore[arg-type]
            case "equation":
                self._render_equation(block)         # type: ignore[arg-type]
            case "list":
                self._render_list(block)             # type: ignore[arg-type]

    # ── Individual renderers ──────────────────────────────────────────────────

    def _render_paragraph(self, block: ParagraphBlock) -> None:
        p = self.doc.add_paragraph(block.text, style=S.BODY_TEXT)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if block.indent:
            p.paragraph_format.first_line_indent = Pt(12)

    def _render_heading(self, block: HeadingBlock) -> None:
        style = {1: S.HEADING_1, 2: S.HEADING_2, 3: S.HEADING_3}[block.level]
        label = ""
        if block.numbering:
            label = f"{block.numbering}. " if block.level == 1 else f"{block.numbering} "
        p = self.doc.add_paragraph(style=style)
        if block.level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}{block.text.upper()}")
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"{label}{block.text}")
            run.italic = True

    def _render_figure(self, block: FigureBlock) -> None:
        img_path = self.storage_base / block.image_ref
        if img_path.exists():
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            fitted = _fit_picture_size(img_path, S.COL_WIDTH_IN, S.MAX_FIGURE_HEIGHT_IN)
            if fitted:
                width_in, height_in = fitted
                run.add_picture(str(img_path), width=Inches(width_in), height=Inches(height_in))
            else:
                run.add_picture(str(img_path), width=Inches(S.COL_WIDTH_IN))
        else:
            p = self.doc.add_paragraph(f"[IMAGE NOT FOUND: {block.image_ref}]", style=S.BODY_TEXT)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cap = self.doc.add_paragraph(style=S.FIGURE_CAPTION)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(f"Fig. {block.anchor.replace('FIG ', '')}. ").bold = True
        cap.add_run(block.caption)

    def _render_wide_figure(self, block: FigureBlock) -> None:
        self._enter_wide()
        img_path = self.storage_base / block.image_ref
        if img_path.exists():
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            fitted = _fit_picture_size(img_path, S.FULL_WIDTH_IN, S.MAX_FIGURE_HEIGHT_IN)
            if fitted:
                width_in, height_in = fitted
                run.add_picture(str(img_path), width=Inches(width_in), height=Inches(height_in))
            else:
                run.add_picture(str(img_path), width=Inches(S.FULL_WIDTH_IN))
        else:
            p = self.doc.add_paragraph(f"[IMAGE NOT FOUND: {block.image_ref}]", style=S.BODY_TEXT)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cap = self.doc.add_paragraph(style=S.FIGURE_CAPTION)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(f"Fig. {block.anchor.replace('FIG ', '')}. ").bold = True
        cap.add_run(block.caption)
        self._exit_wide()

    def _render_table(self, block: TableBlock) -> None:
        if not block.rows:
            return

        # Caption above the table (IEEE convention: table caption before table)
        cap = self.doc.add_paragraph(style=S.FIGURE_CAPTION)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(f"TABLE {block.anchor.replace('TABLE ', '')}").bold = True
        cap.add_run(f"  {block.caption}")

        tbl = self.doc.add_table(rows=len(block.rows), cols=len(block.rows[0]))
        tbl.style = "Normal Table"
        self._apply_table_borders(tbl)
        for r_idx, row_data in enumerate(block.rows):
            row = tbl.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = cell_text
                style = S.TABLE_COL_HEAD if (block.header_row and r_idx == 0) else S.TABLE_COPY
                for para in cell.paragraphs:
                    try:
                        para.style = style
                    except KeyError:
                        pass  # style not in template — leave as-is

    @staticmethod
    def _apply_table_borders(tbl) -> None:
        """Add thin borders to every cell (IEEE standard table look)."""
        from docx.oxml import OxmlElement
        tbl_pr = tbl._tbl.tblPr
        tbl_borders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

    def _render_wide_table(self, block: TableBlock) -> None:
        self._enter_wide()
        self._render_table(block)
        self._exit_wide()

    def _render_equation(self, block: EquationBlock) -> None:
        img_path = self.math_images.get(block.anchor)
        if img_path and img_path.exists():
            para = self.doc.add_paragraph(style=S.EQUATION)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            # Scale equation image: display = half column, inline = auto
            width = Inches(S.COL_WIDTH_IN * 0.8) if not block.inline else Inches(1.2)
            run.add_picture(str(img_path), width=width)
        else:
            # Fallback: plain text LaTeX
            p = self.doc.add_paragraph(style=S.EQUATION)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(block.latex)

        # Equation number (right-aligned)
        eq_num = block.anchor.replace("EQ ", "")
        p = self.doc.add_paragraph(style=S.EQUATION)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"({eq_num})")

    def _render_list(self, block: ListBlock) -> None:
        for item in block.items:
            if block.style == "bullet":
                p = self.doc.add_paragraph(style=S.BULLET_LIST)
            else:
                p = self.doc.add_paragraph(style=S.BODY_TEXT)
                p.paragraph_format.left_indent = Pt(18)
            p.add_run(item)

    # ── References ────────────────────────────────────────────────────────────

    def _render_references(self, refs: list[Reference]) -> None:
        # Section heading
        p = self.doc.add_paragraph(style=S.HEADING_1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("REFERENCES")

        for ref in refs:
            p = self.doc.add_paragraph(style=S.REFERENCES)
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.first_line_indent = Pt(-14)  # hanging indent
            authors_str = ", ".join(ref.authors)
            p.add_run(f"[{ref.key}] ").bold = True
            parts = [authors_str, f'"{ref.title}"']
            if ref.venue:
                parts.append(f"in {ref.venue}")
            if ref.volume:
                parts.append(f"vol. {ref.volume}")
            if ref.issue:
                parts.append(f"no. {ref.issue}")
            if ref.pages:
                parts.append(f"pp. {ref.pages}")
            if ref.year:
                parts.append(str(ref.year))
            if ref.doi:
                parts.append(f"doi: {ref.doi}")
            p.add_run(", ".join(parts) + ".")
